import os
from docx import Document

def check_test_output():
    """
    Check the test output file
    """
    file_path = r"d:\ecommerce procurement\test_output_table.docx"
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    print(f"Checking: test_output_table.docx")
    print("=" * 60)
    
    doc = Document(file_path)
    
    # Count tables
    num_tables = len(doc.tables)
    print(f"\nNumber of tables in document: {num_tables}")
    
    if num_tables > 0:
        print("\nTable details:")
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            print(f"  Table {i+1}: {rows} rows x {cols} columns")
            
            # Show first few cells of first row
            if rows > 0 and cols > 0:
                first_row_text = [table.rows[0].cells[j].text[:20] for j in range(min(5, cols))]
                print(f"    First row: {first_row_text}")
                
                # Show a data row
                if rows > 1:
                    second_row_text = [table.rows[1].cells[j].text[:20] for j in range(min(3, cols))]
                    print(f"    Second row sample: {second_row_text}")
    
    # Search for placeholder text
    print("\nSearching for unreplaced placeholders...")
    found_placeholders = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if ("{{" in text and "}}" in text) or "导入清单" in text or "招标文件表格" in text:
            found_placeholders.append((i, text))
    
    if found_placeholders:
        print(f"\n⚠️ Found {len(found_placeholders)} unreplaced placeholders:")
        for i, text in found_placeholders[:5]:
            print(f"  Para {i}: {text[:100]}")
    else:
        print("\n✅ No unreplaced placeholders found!")

if __name__ == "__main__":
    check_test_output()
