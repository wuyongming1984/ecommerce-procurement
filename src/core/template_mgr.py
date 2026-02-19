import os
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import openpyxl
from openpyxl.styles import Border, Side
from pathlib import Path
import re
from .doc_converter import DocConverter

class TemplateManager:
    """
    Manages loading and rendering of Word templates.
    """
    def __init__(self, template_dir):
        self.template_dir = Path(template_dir)

    def render_template(self, template_path, context, output_path):
        """
        Renders a single Word template with the given context.
        """
        template_filepath = Path(template_path)
        if not template_filepath.exists():
            # If absolute path fails, try relative to template_dir
            template_filepath = self.template_dir / template_path
            if not template_filepath.exists():
                 raise FileNotFoundError(f"Template not found: {template_path}")

        # Convert .doc if needed
        actual_template_path = DocConverter.ensure_docx(str(template_filepath))
        doc = DocxTemplate(actual_template_path)
        
        
        template_vars = doc.get_undeclared_template_variables()
        print(f"DEBUG: Detected variables in template '{template_path}': {template_vars}")

        # Separate table data from other context
        final_context = {}
        table_data_map = {}  # key -> list of dicts
        print(f"DEBUG: Processing context with {len(context)} items.")
        
        for key, value in context.items():
            if isinstance(value, list) and value and isinstance(value[0], dict):
                print(f"DEBUG: Found table data for '{key}' with {len(value)} rows.")
                
                if key not in template_vars:
                    print(f"⚠️ WARNING: Variable '{key}' (Table) NOT found in template '{os.path.basename(str(template_path))}'!")
                    print(f"   Possible Reason: The tag '{{{{ {key} }}}}' might be split by Word formatting (XML tags).")
                    print(f"   Fix: Open the template, Cut the tag '{{{{ {key} }}}}', and Paste it back as 'Keep Text Only'.")
                else:
                    # Use a unique marker that docxtpl will place in the document
                    marker = f"__TABLE_MARKER_{key}__"
                    final_context[key] = marker
                    table_data_map[key] = value
                    print(f"DEBUG: Will insert table for '{key}' via post-processing.")

            elif isinstance(value, str):
                # Check for size parameters: path|width=50|height=30
                parts = value.split('|')
                img_path = parts[0].strip()
                width = None
                height = None
                
                # Parse additional params
                for part in parts[1:]:
                    if part.startswith('width='):
                        val_str = part.split('=')[1]
                        if val_str.lower() == 'full':
                             try:
                                 section = doc.sections[0]
                                 page_width = section.page_width
                                 left_margin = section.left_margin
                                 right_margin = section.right_margin
                                 width = page_width - left_margin - right_margin
                             except:
                                 width = Mm(160)
                        else:
                            try:
                                width = Mm(float(val_str))
                            except: pass
                    elif part.startswith('height='):
                        try:
                            height = Mm(float(part.split('=')[1]))
                        except: pass

                if (img_path.lower().endswith('.png') or \
                    img_path.lower().endswith('.jpg') or \
                    img_path.lower().endswith('.jpeg')):
                    
                    if os.path.exists(img_path):
                        try:
                            if width is None and height is None:
                                width = Mm(40)
                            img = InlineImage(doc, img_path, width=width, height=height)
                            final_context[key] = img
                        except Exception as e:
                           print(f"Error loading image {img_path}: {e}")
                           final_context[key] = img_path
                    else:
                        final_context[key] = value
                else:
                    final_context[key] = value
            else:
                final_context[key] = value

        doc.render(final_context)
        
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        # Post-processing: replace markers with actual tables
        if table_data_map:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
            from copy import deepcopy
            
            out_doc = DocxDocument(output_path)
            
            for key, data in table_data_map.items():
                marker = f"__TABLE_MARKER_{key}__"
                print(f"DEBUG: Post-processing - searching for marker '{marker}'...")
                
                found = False
                for para in out_doc.paragraphs:
                    if marker in para.text:
                        found = True
                        print(f"DEBUG: Found marker for '{key}', inserting table ({len(data)} rows)...")
                        
                        # Build the table
                        headers = list(data[0].keys())
                        table = out_doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'
                        table.autofit = True
                        
                        # Header row
                        for i, header in enumerate(headers):
                            cell = table.rows[0].cells[i]
                            cell.text = str(header)
                            # Bold header
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                        
                        # Data rows
                        for item in data:
                            row_cells = table.add_row().cells
                            for i, header in enumerate(headers):
                                val = item.get(header, "")
                                if val is None: val = ""
                                row_cells[i].text = str(val)
                        
                        # Move table to where the marker paragraph is
                        # by inserting the table XML element right after the marker paragraph
                        para_element = para._element
                        para_element.addnext(table._tbl)
                        
                        # Remove the marker paragraph
                        para_element.getparent().remove(para_element)
                        
                        print(f"✅ Successfully inserted table for '{key}'!")
                        break
                
                if not found:
                    print(f"⚠️ Marker for '{key}' not found in rendered document.")
            
            out_doc.save(output_path)
            print(f"DEBUG: Post-processing complete. Saved to {output_path}")
        
        return output_path

    def _build_table_in_subdoc(self, subdoc, data):
        """
        Helper to build a table in a subdoc from a list of dictionaries.
        """
        if not data:
            return

        headers = list(data[0].keys())
        rows = len(data)
        cols = len(headers)

        # Create table with header row
        table = subdoc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid' # Default style
        table.autofit = True

        # Header
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
        
        # Data
        for item in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                val = item.get(header, "")
                if val is None: val = ""
                row_cells[i].text = str(val)

    def get_template_variables(self, template_path):
        """
        Extracts undeclared variables from the template.
        """
        template_filepath = Path(template_path)
        if not template_filepath.exists():
             template_filepath = self.template_dir / template_path
        
        if not template_filepath.exists():
            return set()

        actual_template_path = DocConverter.ensure_docx(str(template_filepath))
        doc = DocxTemplate(actual_template_path)
        return doc.get_undeclared_template_variables()

    def render_excel_template(self, template_path, context, output_path):
        """
        Renders an Excel template by replacing {{ key }} with values from context.
        """
        template_filepath = Path(template_path)
        if not template_filepath.exists():
             template_filepath = self.template_dir / template_path
        
        if not template_filepath.exists():
            raise FileNotFoundError(f"Excel template not found: {template_path}")

        # Load workbook
        wb = openpyxl.load_workbook(template_filepath)
        
        # Prepare regex for replacement
        pattern = re.compile(r"\{\{\s*(.*?)\s*\}\}")

        # 1. Collect all replacements to avoid modifying while iterating
        replacements = [] # (sheet_idx, row, col, key, original_val)

        for s_idx, sheet in enumerate(wb.worksheets):
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        val = cell.value
                        matches = pattern.findall(val)
                        if matches:
                            # We only handle the first match if it's a table (full cell replacement)
                            # Or multiple matches if just strings.
                            # For simplicity, if a cell contains {{ Table }}, we treat it as a table anchor.
                            for m in matches:
                                key = m.strip()
                                replacements.append({
                                    "sheet_idx": s_idx, 
                                    "row": cell.row, 
                                    "col": cell.column, 
                                    "key": key, 
                                    "val": val, 
                                    "cell": cell
                                })

        # 2. Process replacements (Tables first, usually bottom-up or just handle)
        # We need to distinguish Tables (list) from Variables (str/int)
        
        # Sort by row descending so insertions don't mess up lower rows? 
        # Actually insertions push down, so processing Top-Down requires tracking offset?
        # Processing Bottom-Up is safer for insertions.
        replacements.sort(key=lambda x: (x["sheet_idx"], x["row"]), reverse=True)

        for item in replacements:
            sheet = wb.worksheets[item["sheet_idx"]]
            key = item["key"]
            cell =  sheet.cell(item["row"], item["col"]) # re-acquire cell in case object stale? (Unlikely for exact coords)
            
            if key in context:
                data = context[key]
                
                # Check if it is a list (Table)
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    # It's a table!
                    # Logic: 
                    # 1. Clear the placeholder cell.
                    # 2. Insert N-1 rows (since we use the current row for the first item).
                    #    Or Insert N rows and shift?
                    #    Let's use the current row for the first item.
                    # 3. Fill data.
                    
                    # Columns: The dict keys must map to columns. 
                    # Assumption: The user simply wants the values dumped in order of keys.
                    # Config 'order' determines dict key order.
                    
                    start_row = item["row"]
                    start_col = item["col"]
                    
                    # Clear placeholder
                    clean_val = item["val"].replace(f"{{{{ {key} }}}}", "") # What if other text exists?
                    # If it's a table, we usually expect ONLY the placeholder.
                    # Let's overwrite.
                    
                    count = len(data)
                    if count > 0:
                        # Total rows = 1 (header) + count (data)
                        # Current row is used for header, insert count rows after it
                        if count > 0:
                            sheet.insert_rows(start_row + 1, amount=count)
                        
                        # Fill data
                        keys = list(data[0].keys()) # Order from first dict
                        
                        # Prepare styles
                        thin_border = Border(left=Side(style='thin'), 
                                             right=Side(style='thin'), 
                                             top=Side(style='thin'), 
                                             bottom=Side(style='thin'))
                        from openpyxl.styles import Font, Alignment
                        bold_font = Font(bold=True)
                        center_align = Alignment(horizontal='center', vertical='center')

                        # Write Header Row
                        for j, k in enumerate(keys):
                            header_cell = sheet.cell(start_row, start_col + j)
                            header_cell.value = k
                            header_cell.font = bold_font
                            header_cell.border = thin_border
                            header_cell.alignment = center_align

                        # Write Data Rows (starting from start_row + 1)
                        for i, row_data in enumerate(data):
                            current_r = start_row + 1 + i
                            for j, k in enumerate(keys):
                                target_cell = sheet.cell(current_r, start_col + j)
                                val = row_data.get(k, "")
                                target_cell.value = val
                                target_cell.border = thin_border

                # Standard Variable
                else:
                    # It's a string/number
                    # We might have multiple variables in one string: "Price: {{price}} Unit: {{unit}}"
                    # Re-read current value in case it was modified by another replacement in the same cell?
                    # Since we reverse sorted, and usually one var per cell or multiple,
                    # Safe to just standard replace string.
                    
                    current_val = str(cell.value) if cell.value is not None else ""
                    
                    def replacer(match):
                        k = match.group(1).strip()
                        if k in context:
                            return str(context[k])
                        return match.group(0)
                        
                    new_val_str = pattern.sub(replacer, current_val)
                    
                    # Numeric conversion check
                    if new_val_str != current_val:
                        cell.value = new_val_str
                        try:
                            if "." in new_val_str:
                                cell.value = float(new_val_str)
                            elif new_val_str.isdigit():
                                cell.value = int(new_val_str)
                        except:
                            pass

        # Ensure output directory
        out_dir = Path(output_path).parent
        if not out_dir.exists():
            out_dir.mkdir(parents=True, exist_ok=True)
            
        try:
            wb.save(output_path)
            print(f"DEBUG: Excel saved to {output_path}")
        except Exception as e:
            print(f"ERROR Saving Excel: {e}")
            raise e
            
        return output_path

    def render_and_save(self, template_name, context, output_path):
        if str(template_name).lower().endswith(".xlsx"):
            return self.render_excel_template(template_name, context, output_path)
        return self.render_template(template_name, context, output_path)


class ExcelLoader:
    """
    Handles loading data from Excel files.
    """
    def __init__(self, excel_path):
        self.excel_path = Path(excel_path)

    def load_data(self, sheet_name=None):
        """
        Loads data from the Excel file.
        Returns a list of dictionaries, where keys are column headers.
        """
        if not self.excel_path.exists():
            raise FileNotFoundError(f"Excel file not found: {self.excel_path}")

        wb = openpyxl.load_workbook(self.excel_path, data_only=True)
        if sheet_name:
            ws = wb[sheet_name]
        else:
            ws = wb.active

        data = []
        rows = list(ws.rows)
        if not rows:
            return []

        headers = [cell.value for cell in rows[0]]
        
        for row in rows[1:]:
            row_data = {}
            has_data = False
            for i, cell in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_data[headers[i]] = cell.value
                    if cell.value is not None:
                         has_data = True
            if has_data:
                data.append(row_data)
        
        return data
